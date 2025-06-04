# src/utils.py

import yaml
from pathlib import Path
from docx import Document # type: ignore
from docx.opc.exceptions import PackageNotFoundError

def load_config(config_path_str: str = "config.yaml") -> dict:
    """
    Carrega configurações de um arquivo YAML.

    Args:
        config_path_str: Caminho para o arquivo de configuração YAML.

    Returns:
        Um dicionário com as configurações ou um dicionário vazio se o arquivo não for encontrado
        ou ocorrer um erro.
    """
    config_path = Path(config_path_str)
    if not config_path.is_file():
        # print(f"Info: Arquivo de configuração '{config_path_str}' não encontrado. Usando padrões.")
        return {}
    try:
        with open(config_path, 'r', encoding='utf-8') as f:
            config_data = yaml.safe_load(f)
            return config_data if config_data else {}
    except yaml.YAMLError as e:
        print(f"Erro ao carregar o arquivo de configuração '{config_path_str}': {e}")
        return {}
    except Exception as e:
        print(f"Erro inesperado ao ler o arquivo de configuração '{config_path_str}': {e}")
        return {}

def get_docx_metadata(filepath_str: str) -> dict | None:
    """
    Extrai metadados de um arquivo .docx.

    Args:
        filepath_str: Caminho para o arquivo .docx.

    Returns:
        Um dicionário contendo os metadados ou None se ocorrer um erro.
    """
    filepath = Path(filepath_str)
    if not filepath.is_file():
        print(f"Erro: Arquivo não encontrado em '{filepath_str}'")
        return None
    if filepath.suffix.lower() != ".docx":
        print(f"Erro: O arquivo '{filepath_str}' não é um .docx.")
        return None

    try:
        doc = Document(filepath)
        core_props = doc.core_properties
        
        metadata = {
            "filename": filepath.name,
            "title": core_props.title,
            "author": core_props.author,
            "subject": core_props.subject,
            "keywords": core_props.keywords,
            "category": core_props.category,
            "comments": core_props.comments,
            "last_modified_by": core_props.last_modified_by,
            "created": core_props.created.strftime("%Y-%m-%d %H:%M:%S") if core_props.created else None,
            "modified": core_props.modified.strftime("%Y-%m-%d %H:%M:%S") if core_props.modified else None,
            "revision": core_props.revision,
            "num_paragraphs": len(doc.paragraphs),
            "num_tables": len(doc.tables),
            "num_sections": len(doc.sections),
        }
        # Limpa entradas None ou vazias para uma saída mais limpa
        return {k: v for k, v in metadata.items() if v not in [None, ""]}

    except PackageNotFoundError:
        print(f"Erro: O arquivo '{filepath_str}' não parece ser um arquivo DOCX válido ou está corrompido.")
        return None
    except Exception as e:
        print(f"Erro ao tentar ler metadados de '{filepath_str}': {e}")
        return None

if __name__ == '__main__':
    # Testes rápidos (requer um arquivo .docx de exemplo)
    
    # Criar um arquivo de config de teste
    test_config_content = """
default_output_directory: "meus_pdfs"
custom_setting: 123
"""
    with open("test_config.yaml", "w") as f:
        f.write(test_config_content)
    
    print("--- Testando load_config ---")
    config = load_config("test_config.yaml")
    print(f"Configuração carregada: {config}")
    Path("test_config.yaml").unlink() # Remover arquivo de teste

    config_inexistente = load_config("nao_existe.yaml")
    print(f"Configuração de arquivo inexistente: {config_inexistente}")
    
    # Para testar get_docx_metadata, você precisa de um arquivo .docx
    # Suponha que você tenha um 'example.docx' na raiz do projeto
    example_docx_path = Path("example.docx")
    if example_docx_path.exists():
        print(f"\n--- Testando get_docx_metadata para {example_docx_path} ---")
        metadata = get_docx_metadata(str(example_docx_path))
        if metadata:
            for key, value in metadata.items():
                print(f"  {key}: {value}")
    else:
        print(f"\nArquivo '{example_docx_path}' não encontrado. Pule os testes de metadados.")
